#!/usr/bin/env python3

from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import openpyxl
from docx import Document
from docx.enum.text import WD_BREAK
import os
import tempfile
from datetime import datetime
from werkzeug.utils import secure_filename
import re

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # Change this to a random secret key
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

ALLOWED_EXTENSIONS = {'xlsx', 'xls'}
# Prefer new template when available
NEW_TEMPLATE_DOCX = "New ARN Change form.docx"
OLD_TEMPLATE_DOCX = "Request for Change of Broker.docx"
TEMPLATE_DOCX = NEW_TEMPLATE_DOCX if os.path.exists(NEW_TEMPLATE_DOCX) else OLD_TEMPLATE_DOCX

# Defaults used only when Excel does not provide values
DEFAULT_NEW_ARN_CODE = "310082"
DEFAULT_NEW_ARN_NAME = "Shareway Securities Pvt Ltd"
DEFAULT_EUIN_CODE = "588234"
DEFAULT_EUIN_NAME = "Ajath Anjanappa"
DEFAULT_PLACE = "Bengaluru, Karnataka"


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _looks_like_pan(value: str) -> bool:
    """Detect if a string looks like a PAN number (e.g., ABCDE1234F)."""
    if value is None:
        return False
    s = str(value).strip().upper().replace(" ", "")
    return bool(re.fullmatch(r"[A-Z]{5}[0-9]{4}[A-Z]", s))


def _format_euin(euin_code: str) -> str:
    """Format EUIN code with 'E' prefix."""
    if not euin_code:
        return ""
    # Remove any existing 'E' prefix and add it back
    clean_code = str(euin_code).strip().replace("E", "").replace("e", "")
    return f"E{clean_code}" if clean_code else ""


def read_excel_data(excel_file_path):
    """Read data from Excel file and return as list of dictionaries (one per row).
    Expected columns:
      A: Scheme Name
      B: Folio No
      C: PAN (optional, ignored for scheme name)
      D: Investor [First Holder only]
      E: Old ARN Number
      F: Old ARN Name
    """
    print(f"[DEBUG] Starting Excel file reading: {excel_file_path}")
    workbook = None
    try:
        workbook = openpyxl.load_workbook(excel_file_path, read_only=True)
        sheet = workbook.active
        print(f"[DEBUG] Excel file loaded successfully")
        print(f"[DEBUG] Sheet name: {sheet.title}")
        print(f"[DEBUG] Max row: {sheet.max_row}, Max column: {sheet.max_column}")
        
        data_rows = []
        max_row = sheet.max_row
        
        print(f"[DEBUG] Processing rows 2 to {max_row}")
        
        for row_num in range(2, max_row + 1):
            print(f"[DEBUG] --- Processing Row {row_num} ---")
            
            # Columns
            scheme_a = sheet[f'A{row_num}'].value
            folio_no = sheet[f'B{row_num}'].value
            column_c = sheet[f'C{row_num}'].value
            investor = sheet[f'D{row_num}'].value
            old_arn_number = sheet[f'E{row_num}'].value
            old_arn_name = sheet[f'F{row_num}'].value
            
            print(f"[DEBUG] Raw values: A='{scheme_a}' (type: {type(scheme_a)})")
            print(f"[DEBUG] Raw values: B='{folio_no}' (type: {type(folio_no)})")
            print(f"[DEBUG] Raw values: C='{column_c}' (type: {type(column_c)})")
            print(f"[DEBUG] Raw values: D='{investor}' (type: {type(investor)})")
            print(f"[DEBUG] Raw values: E='{old_arn_number}' (type: {type(old_arn_number)})")
            print(f"[DEBUG] Raw values: F='{old_arn_name}' (type: {type(old_arn_name)})")
            
            # Normalize
            scheme_a_str = str(scheme_a).strip() if scheme_a is not None else ''
            folio_no_str = str(folio_no).strip() if folio_no is not None else ''
            col_c_str = str(column_c).strip() if column_c is not None else ''
            investor_str = str(investor).strip() if investor is not None else ''
            old_arn_number_str = str(old_arn_number).strip() if old_arn_number is not None else ''
            old_arn_name_str = str(old_arn_name).strip() if old_arn_name is not None else ''

            # Detect PAN in column C; if not PAN and non-empty, allow as override for scheme
            pan_from_c = col_c_str.upper().replace(" ", "") if _looks_like_pan(col_c_str) else ''
            scheme_from_c = '' if pan_from_c else col_c_str

            # Final scheme name: prefer Column C when provided and not PAN; else Column A
            scheme_name_str = scheme_from_c if scheme_from_c else scheme_a_str
            # PAN for backward compatibility
            pan_str = pan_from_c
            
            print(f"[DEBUG] Processed: SCHEME_A='{scheme_a_str}', SCHEME_C='{scheme_from_c}', SCHEME_FINAL='{scheme_name_str}', PAN_FROM_C='{pan_str}'")
            print(f"[DEBUG] Processed: FOLIO='{folio_no_str}', INVESTOR='{investor_str}'")
            print(f"[DEBUG] Processed: OLD_ARN_NUM='{old_arn_number_str}', OLD_ARN_NAME='{old_arn_name_str}'")
            
            has_data = any([
                scheme_name_str and scheme_name_str != 'None',
                folio_no_str and folio_no_str != 'None',
                investor_str and investor_str != 'None'
            ])
            print(f"[DEBUG] Row {row_num} has data: {has_data}")
            if not has_data:
                print(f"[DEBUG] SKIPPING empty row {row_num}")
                continue
                
            data = {
                # Use a generic header indicator for the new template
                'mutual_fund': 'Multiple',
                'folio_no': folio_no_str,
                'scheme_name': scheme_name_str,
                'investor': investor_str,
                'pan': pan_str,
                # Old ARN details from Excel
                'old_arn_code': old_arn_number_str,
                'old_arn_name': old_arn_name_str,
                # Hardcoded new ARN values
                'new_arn_code': DEFAULT_NEW_ARN_CODE,
                'new_arn_name': DEFAULT_NEW_ARN_NAME,
                'new_sub_arn_code': '',
                'new_euin_code': _format_euin(DEFAULT_EUIN_CODE),
                'sub_distributor_name': '',
                'euin_name': DEFAULT_EUIN_NAME,
                'arn_euin_holder_signature': '',
                'new_distributor_staff_info': '',
                'place': DEFAULT_PLACE,
            }
            data_rows.append(data)
            print(f"[DEBUG] ADDED row {row_num} to data_rows (total now: {len(data_rows)})")
        
        print(f"[DEBUG] Excel reading complete. Total data rows found: {len(data_rows)}")
        return data_rows if data_rows else None
    except Exception as e:
        print(f"[DEBUG] ERROR reading Excel file: {str(e)}")
        return None
    finally:
        if workbook:
            workbook.close()
            print(f"[DEBUG] Excel workbook closed")


def populate_single_page_old_form(doc, data):
    """Populate the legacy 'Request for Change of Broker.docx' template."""
    print(f"[DEBUG] Starting to populate legacy single page with data: {data}")
    print(f"[DEBUG] Document has {len(doc.paragraphs)} paragraphs")
    
    fields_populated = 0
    
    # Process paragraphs with precise formatting preservation
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        
        # Handle "Mutual Fund: " line (Paragraph 3)
        if original_text.strip() == 'Mutual Fund:':
            print(f"[DEBUG] Found Mutual Fund field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("  Mutual Fund: ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['mutual_fund']))
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Mutual Fund: '{data['mutual_fund']}'")
        
        # Handle "Folio No:* ... PAN:* " line (Paragraph 4)
        elif 'Folio No:*' in original_text and 'PAN:*' in original_text:
            print(f"[DEBUG] Found Folio/PAN field at paragraph {i}")
            paragraph.clear()
            # Add bold "Folio No:*" label
            run1 = paragraph.add_run("      Folio No:* ")
            run1.bold = True
            # Add underlined folio number
            run2 = paragraph.add_run(str(data['folio_no']))
            run2.underline = True
            # Add spacing
            paragraph.add_run("                                                                                                          ")
            # Add bold "PAN:*" label
            run3 = paragraph.add_run("PAN:* ")
            run3.bold = True
            # Add underlined PAN
            run4 = paragraph.add_run(str(data['pan']))
            run4.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Folio: '{data['folio_no']}', PAN: '{data['pan']}'")
        
        # Handle "Investor [First Holder only]:  " line (Paragraph 5)
        elif original_text.strip() == 'Investor [First Holder only]:':
            print(f"[DEBUG] Found Investor field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("  Investor [First Holder only]: ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['investor']).strip())
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Investor: '{data['investor']}'")
        
        # Handle acknowledgement slip fields
        elif original_text.strip() == 'Mutual Fund :':
            print(f"[DEBUG] Found Acknowledgement Mutual Fund field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("Mutual Fund : ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['mutual_fund']))
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Ack Mutual Fund: '{data['mutual_fund']}'")
        elif 'Folio No :' in original_text and 'Date of Receipt:' in original_text:
            print(f"[DEBUG] Found Acknowledgement Folio field at paragraph {i}")
            paragraph.clear()
            # Add bold "Folio No :" label
            run1 = paragraph.add_run("Folio No : ")
            run1.bold = True
            # Add underlined folio number
            run2 = paragraph.add_run(str(data['folio_no']))
            run2.underline = True
            # Add spacing and Date of Receipt
            paragraph.add_run("                               \t\t                                       Date of Receipt:\t")
            fields_populated += 1
            print(f"[DEBUG] Populated Ack Folio: '{data['folio_no']}'")
    
    print(f"[DEBUG] Legacy single page population complete. Fields populated: {fields_populated}")


def _replace_text_anywhere(doc, replacements):
    """Replace text in all text nodes, including inside shapes/textboxes.
    replacements: dict where key can be a string token or a tuple/list of token variants.
    Only replaces when the text node matches the token exactly (ignoring surrounding whitespace),
    to avoid duplicating values when the field is already filled.
    """
    text_nodes = doc.part.element.xpath('.//w:t')
    replaced_counts = {}

    def ensure_iterable(key):
        if isinstance(key, (list, tuple)):
            return list(key)
        return [key]

    for t in text_nodes:
        current = t.text or ''
        current_norm = current.replace('\xa0', ' ').strip()
        for key, val in replacements.items():
            tokens = ensure_iterable(key)
            for token in tokens:
                token_norm = str(token).strip()
                if current_norm == token_norm:
                    # Special handling for EUIN to avoid double "E"
                    if 'EUIN No.: E' in token_norm:
                        # Replace the entire "EUIN No.: E" with "EUIN No.: E588234"
                        t.text = f"EUIN No.: {val}"
                    else:
                        # Normal replacement
                        t.text = f"{token_norm} {val}"
                    replaced_counts[token_norm] = replaced_counts.get(token_norm, 0) + 1
    print(f"[DEBUG] Textbox replacements: {replaced_counts}")


def populate_single_page_new_form(doc, data):
    """Populate the new 'New ARN Change form.docx' template for a single-entry page (legacy mode)."""
    print(f"[DEBUG] Starting to populate NEW single page with data: {data}")
    print(f"[DEBUG] Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")

    # Paragraph fills (header/date/place only)
    today_str = data.get('date') or datetime.now().strftime('%d-%m-%Y')
    place_str = data.get('place', '')

    for i, paragraph in enumerate(doc.paragraphs):
        txt = paragraph.text.strip()
        # Header line with Mutual Fund and Date
        if ('Mutual Fund' in txt) and ('Date' in txt):
            print(f"[DEBUG] Found MF/Date header at paragraph {i}")
            paragraph.clear()
            paragraph.add_run(f"{data.get('mutual_fund', '')} Mutual Fund\t\t\t\tDate: {today_str}")
        # Footer Date
        elif txt.startswith('Date:') and 'Mutual Fund' not in txt:
            print(f"[DEBUG] Found Date footer at paragraph {i}")
            paragraph.clear()
            paragraph.add_run(f"Date: {today_str}")
        # Footer Place
        elif txt.lower().startswith('place'):
            print(f"[DEBUG] Found Place footer at paragraph {i}")
            paragraph.clear()
            paragraph.add_run(f"Place: {place_str}")

    # Replace text (left and right) using safe token replacements
    _replace_text_anywhere(doc, {
        ('New ARN-.', 'New ARN:', 'New ARN -'): data.get('new_arn_code', DEFAULT_NEW_ARN_CODE),
        ("Sub-Distributor's ARN", "Sub-Distributor's ARN"): data.get('new_sub_arn_code', ''),
        ('EUIN No.: E', 'EUIN No.:', 'EUIN No:', 'EUIN No', 'EUIN'): data.get('new_euin_code', _format_euin(DEFAULT_EUIN_CODE)),
        'ARN Name:': data.get('new_arn_name', DEFAULT_NEW_ARN_NAME),
        ("Sub-Distributor's name :", "Sub-Distributor's name :"): data.get('sub_distributor_name', ''),
        'EUIN Name:': data.get('euin_name', ''),
        ('Signature of ARN/EUIN Holder:', 'Signature of ARN/ EUIN Holder:'): data.get('arn_euin_holder_signature', ''),
        (
            'Name, Designation, Employee code of new distributor (if non individual)',
            'Name, Designation, Employee code of new distributor'
        ): data.get('new_distributor_staff_info', ''),
    })

    # Table fills
    tables = doc.tables
    # Table 0: Folio/Scheme list (single entry mode)
    if len(tables) >= 1 and len(tables[0].rows) >= 2:
        try:
            tables[0].cell(1, 0).text = str(data.get('folio_no', '')).strip()
            tables[0].cell(1, 1).text = str(data.get('scheme_name', '')).strip()
            print("[DEBUG] Filled Table 0: Folio and Scheme")
        except Exception as e:
            print(f"[DEBUG] Could not fill Table 0: {e}")

    # Table 1: ARN details
    if len(tables) >= 2 and len(tables[1].rows) >= 2 and len(tables[1].rows[1].cells) >= 6:
        try:
            old_arn_code = str(data.get('old_arn_code', '')).strip()
            old_arn_name = str(data.get('old_arn_name', '')).strip()
            new_arn_code = str(data.get('new_arn_code', DEFAULT_NEW_ARN_CODE)).strip()
            new_arn_name = str(data.get('new_arn_name', DEFAULT_NEW_ARN_NAME)).strip()
            new_sub_arn = str(data.get('new_sub_arn_code', '')).strip()
            new_euin = str(data.get('new_euin_code', _format_euin(DEFAULT_EUIN_CODE))).strip()

            row = tables[1].rows[1]
            row.cells[0].text = old_arn_code
            row.cells[1].text = old_arn_name
            row.cells[2].text = new_arn_code
            row.cells[3].text = new_arn_name
            row.cells[4].text = new_sub_arn
            row.cells[5].text = new_euin
            print("[DEBUG] Filled Table 1: ARN block")
        except Exception as e:
            print(f"[DEBUG] Could not fill Table 1: {e}")

    # Table 2: Investor details (names only, signatures left blank)
    if len(tables) >= 3 and len(tables[2].rows) >= 3:
        try:
            # Row 1 is headers, Row 2 is "Name" row
            name_row = tables[2].rows[1]
            first_holder = str(data.get('investor', '')).strip()
            second_holder = str(data.get('second_holder', '')).strip()
            third_holder = str(data.get('third_holder', '')).strip()
            # Columns: [label, 1st, 2nd, 3rd]
            if len(name_row.cells) >= 4:
                name_row.cells[1].text = first_holder
                name_row.cells[2].text = second_holder
                name_row.cells[3].text = third_holder
                print("[DEBUG] Filled Table 2: Investor names")
        except Exception as e:
            print(f"[DEBUG] Could not fill Table 2: {e}")

    print("[DEBUG] NEW single page population complete")


def populate_single_page_new_form_chunk(doc, data_chunk):
    """Populate the new template with up to 6 rows on a single page."""
    print(f"[DEBUG] Populating NEW template page with {len(data_chunk)} row(s)")
    # Use first row for shared fields
    first = data_chunk[0]
    # Paragraphs: header/date/place
    today_str = first.get('date') or datetime.now().strftime('%d-%m-%Y')
    place_str = first.get('place', DEFAULT_PLACE)

    # If all mutual funds are the same, use it; else mark as Multiple
    mf_values = {d.get('mutual_fund','') for d in data_chunk}
    header_mf = list(mf_values)[0] if len(mf_values) == 1 else 'Multiple'

    for i, paragraph in enumerate(doc.paragraphs):
        txt = paragraph.text.strip()
        if ('Mutual Fund' in txt) and ('Date' in txt):
            paragraph.clear()
            paragraph.add_run(f"{header_mf} Mutual Fund\t\t\t\tDate: {today_str}")
        elif txt.startswith('Date:') and 'Mutual Fund' not in txt:
            paragraph.clear()
            paragraph.add_run(f"Date: {today_str}")
        elif txt.lower().startswith('place'):
            paragraph.clear()
            paragraph.add_run(f"Place: {place_str}")

    # Replace text (left and right) using safe token replacements
    _replace_text_anywhere(doc, {
        ('New ARN-.', 'New ARN:', 'New ARN -'): first.get('new_arn_code', DEFAULT_NEW_ARN_CODE),
        ("Sub-Distributor's ARN", "Sub-Distributor's ARN"): first.get('new_sub_arn_code', ''),
        ('EUIN No.: E', 'EUIN No.:', 'EUIN No:', 'EUIN No', 'EUIN'): first.get('new_euin_code', _format_euin(DEFAULT_EUIN_CODE)),
        'ARN Name:': first.get('new_arn_name', DEFAULT_NEW_ARN_NAME),
        ("Sub-Distributor's name :", "Sub-Distributor's name :"): first.get('sub_distributor_name', ''),
        'EUIN Name:': first.get('euin_name', ''),
        ('Signature of ARN/EUIN Holder:', 'Signature of ARN/ EUIN Holder:'): first.get('arn_euin_holder_signature', ''),
        (
            'Name, Designation, Employee code of new distributor (if non individual)',
            'Name, Designation, Employee code of new distributor'
        ): first.get('new_distributor_staff_info', ''),
    })

    # Tables
    tables = doc.tables
    # Table 0: fill up to 6 rows
    if len(tables) >= 1:
        t0 = tables[0]
        max_fill = min(6, len(data_chunk))
        for i in range(max_fill):
            try:
                t0.cell(i+1, 0).text = str(data_chunk[i].get('folio_no',''))
                t0.cell(i+1, 1).text = str(data_chunk[i].get('scheme_name',''))
            except Exception as e:
                print(f"[DEBUG] Could not fill Table 0 row {i+1}: {e}")

    # Table 1: use first row values
    if len(tables) >= 2 and len(tables[1].rows) >= 2 and len(tables[1].rows[1].cells) >= 6:
        try:
            row = tables[1].rows[1]
            row.cells[0].text = str(first.get('old_arn_code',''))
            row.cells[1].text = str(first.get('old_arn_name',''))
            row.cells[2].text = str(first.get('new_arn_code', DEFAULT_NEW_ARN_CODE))
            row.cells[3].text = str(first.get('new_arn_name', DEFAULT_NEW_ARN_NAME))
            row.cells[4].text = str(first.get('new_sub_arn_code',''))
            row.cells[5].text = str(first.get('new_euin_code', _format_euin(DEFAULT_EUIN_CODE)))
        except Exception as e:
            print(f"[DEBUG] Could not fill Table 1: {e}")

    # Table 2: investor names (from first row)
    if len(tables) >= 3 and len(tables[2].rows) >= 3:
        try:
            name_row = tables[2].rows[1]
            if len(name_row.cells) >= 4:
                name_row.cells[1].text = str(first.get('investor',''))
                name_row.cells[2].text = str(first.get('second_holder',''))
                name_row.cells[3].text = str(first.get('third_holder',''))
        except Exception as e:
            print(f"[DEBUG] Could not fill Table 2: {e}")

    print("[DEBUG] NEW chunk page population complete")


def populate_single_page_auto(doc, data):
    """Detect template type and populate accordingly."""
    try:
        # Heuristic: the new form has 3 tables and a header with 'Mutual Fund' and 'Date:' on same line
        has_new_signature = (
            len(doc.tables) >= 3 or any(('Mutual Fund' in p.text and 'Date' in p.text) for p in doc.paragraphs)
        )
        if has_new_signature and os.path.basename(TEMPLATE_DOCX) == NEW_TEMPLATE_DOCX:
            return populate_single_page_new_form(doc, data)
        # Fallback to old form logic
        return populate_single_page_old_form(doc, data)
    except Exception as e:
        print(f"[DEBUG] Auto population error, falling back to old form: {e}")
        return populate_single_page_old_form(doc, data)


def chunk_list(items, size):
    return [items[i:i+size] for i in range(0, len(items), size)]


def populate_word_document(template_path, data_list, output_path):
    """Populate Word document with multiple pages of Excel data"""
    print(f"[DEBUG] Starting Word document population")
    print(f"[DEBUG] Template path: {template_path}")
    print(f"[DEBUG] Output path: {output_path}")
    print(f"[DEBUG] Data list contains {len(data_list)} entries")
    
    try:
        # If using new template, group 6 rows per page
        if os.path.basename(template_path) == NEW_TEMPLATE_DOCX:
            print("[DEBUG] New template detected - grouping 6 rows per page")
            pages = chunk_list(data_list, 6)
            print(f"[DEBUG] Total pages after grouping: {len(pages)}")

            # Load and populate first page
            output_doc = Document(template_path)
            populate_single_page_new_form_chunk(output_doc, pages[0])

            # Remaining pages
            for page_index in range(1, len(pages)):
                # Add page break
                if output_doc.paragraphs:
                    last_para = output_doc.paragraphs[-1]
                    run = last_para.add_run()
                    run.add_break(WD_BREAK.PAGE)
                else:
                    break_para = output_doc.add_paragraph()
                    run = break_para.add_run()
                    run.add_break(WD_BREAK.PAGE)

                # Populate fresh template and copy into output
                template_doc = Document(template_path)
                populate_single_page_new_form_chunk(template_doc, pages[page_index])

                for element in template_doc.element.body:
                    element_type = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                    if element_type == 'sectPr':
                        continue
                    output_doc.element.body.append(element)

            output_doc.save(output_path)
            print(f"[DEBUG] New-template multi-page document saved successfully")
            return len(pages)

        # Legacy handling (one row per page)
        # For single page, use simpler approach
        if len(data_list) == 1:
            print(f"[DEBUG] Single page mode - using direct template modification")
            doc = Document(template_path)
            print(f"[DEBUG] Template loaded with {len(doc.paragraphs)} paragraphs")
            populate_single_page_auto(doc, data_list[0])
            doc.save(output_path)
            print(f"[DEBUG] Single page document saved successfully")
            return 1
        
        print(f"[DEBUG] Multi-page mode - using first page as base")
        # For multiple pages, start with the first populated template as base
        print(f"[DEBUG] === Processing Page 1 of {len(data_list)} ===")
        print(f"[DEBUG] Page data: {data_list[0]}")
        
        # Load and populate the first page as the base document
        output_doc = Document(template_path)
        print(f"[DEBUG] Base template loaded with {len(output_doc.paragraphs)} paragraphs")
        populate_single_page_auto(output_doc, data_list[0])
        print(f"[DEBUG] Base template populated for page 1")
        
        # Process remaining pages (if any)
        for page_index in range(1, len(data_list)):
            data = data_list[page_index]
            print(f"[DEBUG] === Processing Page {page_index + 1} of {len(data_list)} ===")
            print(f"[DEBUG] Page data: {data}")
            
            # Add page break to the last paragraph of the current document
            print(f"[DEBUG] Adding page break to last paragraph before page {page_index + 1}")
            if output_doc.paragraphs:
                last_para = output_doc.paragraphs[-1]
                run = last_para.add_run()
                run.add_break(WD_BREAK.PAGE)
                print(f"[DEBUG] Page break added to existing paragraph")
            else:
                # Fallback if no paragraphs exist
                break_para = output_doc.add_paragraph()
                run = break_para.add_run()
                run.add_break(WD_BREAK.PAGE)
                print(f"[DEBUG] Page break added as new paragraph")
            
            # Load a fresh template for this page
            template_doc = Document(template_path)
            print(f"[DEBUG] Fresh template loaded with {len(template_doc.paragraphs)} paragraphs")
            
            # Populate this template with the current row's data
            populate_single_page_auto(template_doc, data)
            print(f"[DEBUG] Template populated for page {page_index + 1}")
            
            # Copy all paragraphs and elements from template to output (except sectPr)
            elements_copied = 0
            for element in template_doc.element.body:
                element_type = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                if element_type == 'sectPr':
                    continue
                output_doc.element.body.append(element)
                elements_copied += 1
            print(f"[DEBUG] Copied {elements_copied} elements from template into output")
        
        # Save the multi-page document
        output_doc.save(output_path)
        print(f"[DEBUG] Multi-page document saved successfully")
        return len(data_list)  # Return number of pages created
    except Exception as e:
        print(f"[DEBUG] ERROR populating Word document: {str(e)}")
        import traceback
        print(f"[DEBUG] Full traceback: {traceback.format_exc()}")
        return False


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file selected')
        return redirect(url_for('index'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected')
        return redirect(url_for('index'))
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        temp_excel_path = None
        
        try:
            # Create temporary file for uploaded Excel
            temp_excel_fd, temp_excel_path = tempfile.mkstemp(suffix='.xlsx')
            
            # Close the file descriptor and save the uploaded file
            os.close(temp_excel_fd)
            file.save(temp_excel_path)
            
            # Read data from Excel (now returns list of dictionaries)
            print(f"[DEBUG] About to read Excel data from: {temp_excel_path}")
            excel_data = read_excel_data(temp_excel_path)
            
            print(f"[DEBUG] Excel reading result: {excel_data}")
            if excel_data is None or len(excel_data) == 0:
                print(f"[DEBUG] No data found in Excel file")
                flash('Error reading Excel file or no data found. Please check the file format.')
                return redirect(url_for('index'))
            
            print(f"[DEBUG] Successfully read {len(excel_data)} data rows from Excel")
            
            # Create output file with page count
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Page count will be computed inside populate_word_document for new template
            output_filename = f"Populated_ARN_Form_{timestamp}.docx"
            output_path = os.path.join(tempfile.gettempdir(), output_filename)
            
            print(f"[DEBUG] Will create output file: {output_filename}")
            print(f"[DEBUG] Full output path: {output_path}")
            
            # Populate Word document
            print(f"[DEBUG] About to populate Word document using template '{TEMPLATE_DOCX}'")
            result = populate_word_document(TEMPLATE_DOCX, excel_data, output_path)
            print(f"[DEBUG] Word document population result: {result}")
            
            if result:
                print(f"[DEBUG] Successfully created document, sending to user")
                return send_file(output_path, as_attachment=True, 
                               download_name=output_filename,
                               mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            else:
                print(f"[DEBUG] Failed to create document")
                flash('Error processing the document. Please try again.')
                return redirect(url_for('index'))
                
        except Exception as e:
            flash(f'Error processing file: {str(e)}')
            return redirect(url_for('index'))
        finally:
            # Clean up temporary Excel file
            if temp_excel_path and os.path.exists(temp_excel_path):
                try:
                    os.unlink(temp_excel_path)
                except PermissionError:
                    pass
    else:
        flash('Please upload a valid Excel file (.xlsx or .xls)')
        return redirect(url_for('index'))


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8000)