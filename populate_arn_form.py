#!/usr/bin/env python3

import openpyxl
from docx import Document
from docx.enum.text import WD_BREAK
import os
from datetime import datetime

def read_excel_data(excel_file_path):
    """Read data from Excel file and return as list of dictionaries (one per row)"""
    print(f"[DEBUG] Starting Excel file reading: {excel_file_path}")
    workbook = None
    try:
        workbook = openpyxl.load_workbook(excel_file_path)
        sheet = workbook.active
        print(f"[DEBUG] Excel file loaded successfully")
        print(f"[DEBUG] Sheet name: {sheet.title}")
        print(f"[DEBUG] Max row: {sheet.max_row}, Max column: {sheet.max_column}")
        
        # Get all data rows (row 1 contains headers, data starts from row 2)
        data_rows = []
        max_row = sheet.max_row
        
        print(f"[DEBUG] Processing rows 2 to {max_row}")
        
        for row_num in range(2, max_row + 1):  # Start from row 2, go to last row
            print(f"[DEBUG] --- Processing Row {row_num} ---")
            
            # Check if row has any data (skip completely empty rows)
            mutual_fund = sheet[f'A{row_num}'].value
            folio_no = sheet[f'B{row_num}'].value
            pan = sheet[f'C{row_num}'].value
            investor = sheet[f'D{row_num}'].value
            
            print(f"[DEBUG] Raw values: MF='{mutual_fund}' (type: {type(mutual_fund)})")
            print(f"[DEBUG] Raw values: FN='{folio_no}' (type: {type(folio_no)})")
            print(f"[DEBUG] Raw values: PAN='{pan}' (type: {type(pan)})")
            print(f"[DEBUG] Raw values: INV='{investor}' (type: {type(investor)})")
            
            # Convert to strings and strip whitespace for better empty detection
            mutual_fund_str = str(mutual_fund).strip() if mutual_fund is not None else ''
            folio_no_str = str(folio_no).strip() if folio_no is not None else ''
            pan_str = str(pan).strip() if pan is not None else ''
            investor_str = str(investor).strip() if investor is not None else ''
            
            print(f"[DEBUG] Processed values: MF='{mutual_fund_str}'")
            print(f"[DEBUG] Processed values: FN='{folio_no_str}'")
            print(f"[DEBUG] Processed values: PAN='{pan_str}'")
            print(f"[DEBUG] Processed values: INV='{investor_str}'")
            
            # Skip row if all cells are empty or contain only 'None'
            has_data = any([mutual_fund_str and mutual_fund_str != 'None', 
                           folio_no_str and folio_no_str != 'None',
                           pan_str and pan_str != 'None', 
                           investor_str and investor_str != 'None'])
            
            print(f"[DEBUG] Row {row_num} has data: {has_data}")
            
            if not has_data:
                print(f"[DEBUG] SKIPPING empty row {row_num}")
                continue
                
            data = {
                'mutual_fund': mutual_fund_str,
                'folio_no': folio_no_str,
                'pan': pan_str,
                'investor': investor_str
            }
            data_rows.append(data)
            print(f"[DEBUG] ADDED row {row_num} to data_rows (total now: {len(data_rows)})")
        
        print(f"[DEBUG] Excel reading complete. Total data rows found: {len(data_rows)}")
        return data_rows if data_rows else None
    except Exception as e:
        print(f"[DEBUG] ERROR reading Excel file: {str(e)}")
        return None
    finally:
        # Ensure workbook is properly closed
        if workbook:
            workbook.close()
            print(f"[DEBUG] Excel workbook closed")

def populate_single_page(doc, data):
    """Helper function to populate a single page with data"""
    print(f"[DEBUG] Starting to populate single page with data: {data}")
    print(f"[DEBUG] Document has {len(doc.paragraphs)} paragraphs")
    
    fields_populated = 0
    
    # Process paragraphs with precise formatting preservation
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        
        # Handle "Mutual Fund: " line (Paragraph 3)
        if original_text.strip() == 'Mutual Fund:':
            print(f"[DEBUG] Found Mutual Fund field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("  Mutual Fund: ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['mutual_fund']))
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Mutual Fund: '{data['mutual_fund']}'")
        
        # Handle "Folio No:* ... PAN:* " line (Paragraph 4)
        elif 'Folio No:*' in original_text and 'PAN:*' in original_text:
            print(f"[DEBUG] Found Folio/PAN field at paragraph {i}")
            paragraph.clear()
            # Add bold "Folio No:*" label
            run1 = paragraph.add_run("      Folio No:* ")
            run1.bold = True
            # Add underlined folio number
            run2 = paragraph.add_run(str(data['folio_no']))
            run2.underline = True
            # Add spacing
            paragraph.add_run("                                                                                                          ")
            # Add bold "PAN:*" label
            run3 = paragraph.add_run("PAN:* ")
            run3.bold = True
            # Add underlined PAN
            run4 = paragraph.add_run(str(data['pan']))
            run4.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Folio: '{data['folio_no']}', PAN: '{data['pan']}'")
        
        # Handle "Investor [First Holder only]:  " line (Paragraph 5)
        elif original_text.strip() == 'Investor [First Holder only]:':
            print(f"[DEBUG] Found Investor field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("  Investor [First Holder only]: ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['investor']).strip())
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Investor: '{data['investor']}'")
        
        # Handle acknowledgement slip fields
        elif original_text.strip() == 'Mutual Fund :':
            print(f"[DEBUG] Found Acknowledgement Mutual Fund field at paragraph {i}")
            paragraph.clear()
            # Add bold label
            run1 = paragraph.add_run("Mutual Fund : ")
            run1.bold = True
            # Add underlined value
            run2 = paragraph.add_run(str(data['mutual_fund']))
            run2.underline = True
            fields_populated += 1
            print(f"[DEBUG] Populated Ack Mutual Fund: '{data['mutual_fund']}'")
        elif 'Folio No :' in original_text and 'Date of Receipt:' in original_text:
            print(f"[DEBUG] Found Acknowledgement Folio field at paragraph {i}")
            paragraph.clear()
            # Add bold "Folio No :" label
            run1 = paragraph.add_run("Folio No : ")
            run1.bold = True
            # Add underlined folio number
            run2 = paragraph.add_run(str(data['folio_no']))
            run2.underline = True
            # Add spacing and Date of Receipt
            paragraph.add_run("                              		                                       Date of Receipt:	")
            fields_populated += 1
            print(f"[DEBUG] Populated Ack Folio: '{data['folio_no']}'")
    
    print(f"[DEBUG] Single page population complete. Fields populated: {fields_populated}")

def populate_word_document(template_path, data_list, output_path):
    """Populate Word document with multiple pages of Excel data"""
    print(f"[DEBUG] Starting Word document population")
    print(f"[DEBUG] Template path: {template_path}")
    print(f"[DEBUG] Output path: {output_path}")
    print(f"[DEBUG] Data list contains {len(data_list)} entries")
    
    try:
        # For single page, use simpler approach
        if len(data_list) == 1:
            print(f"[DEBUG] Single page mode - using direct template modification")
            doc = Document(template_path)
            print(f"[DEBUG] Template loaded with {len(doc.paragraphs)} paragraphs")
            populate_single_page(doc, data_list[0])
            doc.save(output_path)
            print(f"[DEBUG] Single page document saved successfully")
            return 1
        
        print(f"[DEBUG] Multi-page mode - using first page as base")
        # For multiple pages, start with the first populated template as base
        print(f"[DEBUG] === Processing Page 1 of {len(data_list)} ===")
        print(f"[DEBUG] Page data: {data_list[0]}")
        
        # Load and populate the first page as the base document
        output_doc = Document(template_path)
        print(f"[DEBUG] Base template loaded with {len(output_doc.paragraphs)} paragraphs")
        populate_single_page(output_doc, data_list[0])
        print(f"[DEBUG] Base template populated for page 1")
        
        # Process remaining pages (if any)
        for page_index in range(1, len(data_list)):
            data = data_list[page_index]
            print(f"[DEBUG] === Processing Page {page_index + 1} of {len(data_list)} ===")
            print(f"[DEBUG] Page data: {data}")
            
            # Add page break to the last paragraph of the current document
            print(f"[DEBUG] Adding page break to last paragraph before page {page_index + 1}")
            if output_doc.paragraphs:
                last_para = output_doc.paragraphs[-1]
                run = last_para.add_run()
                run.add_break(WD_BREAK.PAGE)
                print(f"[DEBUG] Page break added to existing paragraph")
            else:
                # Fallback if no paragraphs exist
                break_para = output_doc.add_paragraph()
                run = break_para.add_run()
                run.add_break(WD_BREAK.PAGE)
                print(f"[DEBUG] Page break added as new paragraph")
            
            # Load a fresh template for this page
            template_doc = Document(template_path)
            print(f"[DEBUG] Fresh template loaded with {len(template_doc.paragraphs)} paragraphs")
            
            # Populate this template with the current row's data
            populate_single_page(template_doc, data)
            print(f"[DEBUG] Template populated for page {page_index + 1}")
            
            # Count elements before copying
            elements_before = len(output_doc.element.body)
            print(f"[DEBUG] Output document has {elements_before} elements before copying")
            
            # Copy all paragraphs and elements from template to output (except sectPr)
            elements_copied = 0
            for element in template_doc.element.body:
                element_type = element.tag.split('}')[-1] if '}' in element.tag else element.tag
                print(f"[DEBUG] Processing element type: {element_type}")
                
                # Skip sectPr (section properties) elements to avoid blank pages
                if element_type == 'sectPr':
                    print(f"[DEBUG] SKIPPING sectPr element to prevent blank page")
                    continue
                
                output_doc.element.body.append(element)
                elements_copied += 1
                print(f"[DEBUG] Copied element {elements_copied}: {element_type}")
            
            elements_after = len(output_doc.element.body)
            print(f"[DEBUG] Copied {elements_copied} elements from template")
            print(f"[DEBUG] Output document now has {elements_after} elements")
            print(f"[DEBUG] Output document now has {len(output_doc.paragraphs)} paragraphs")
        
        print(f"[DEBUG] Final document has {len(output_doc.paragraphs)} paragraphs")
        print(f"[DEBUG] Final document has {len(output_doc.element.body)} body elements")
        
        # Save the multi-page document
        output_doc.save(output_path)
        print(f"[DEBUG] Multi-page document saved successfully")
        return len(data_list)  # Return number of pages created
    except Exception as e:
        print(f"[DEBUG] ERROR populating Word document: {str(e)}")
        import traceback
        print(f"[DEBUG] Full traceback: {traceback.format_exc()}")
        return False

def main():
    # File paths
    excel_file = "Format for ARN change.xlsx"
    docx_file = "Request for Change of Broker.docx"
    
    # Check if files exist
    if not os.path.exists(excel_file):
        print(f"Error: Excel file '{excel_file}' not found!")
        return
    
    if not os.path.exists(docx_file):
        print(f"Error: Word document '{docx_file}' not found!")
        return
    
    try:
        # Read data from Excel (now returns list of dictionaries)
        print("Reading data from Excel file...")
        excel_data = read_excel_data(excel_file)
        
        if excel_data is None or len(excel_data) == 0:
            print("Error: No data found in Excel file or file is empty.")
            return
        
        print(f"Excel data loaded - {len(excel_data)} row(s) found:")
        for i, data in enumerate(excel_data, 1):
            print(f"\nRow {i}:")
            for key, value in data.items():
                print(f"  {key}: {value}")
        
        # Create output file with page count
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        page_count = len(excel_data)
        output_file = f"Populated_ARN_Form_{page_count}pages_{timestamp}.docx"
        
        # Populate Word document (now handles multiple pages)
        print(f"\nPopulating Word document with {page_count} page(s)...")
        result = populate_word_document(docx_file, excel_data, output_file)
        
        if result:
            print(f"\nSuccess! Generated {result} page(s) from {len(excel_data)} Excel row(s).")
            print(f"Output file: {output_file}")
        else:
            print("Error: Failed to populate Word document.")
        
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main()